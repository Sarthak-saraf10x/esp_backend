import os
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import re

def generate_docx(title, content, output_dir="documents"):
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{title.replace(' ', '_')}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    doc.add_heading(title, 0)
    
    paragraphs = content.split('\n')
    for p in paragraphs:
        clean_p = p.strip()
        if not clean_p:
            continue
            
        if clean_p.startswith('# '):
            doc.add_heading(clean_p[2:].replace('**', '').replace('*', ''), level=1)
        elif clean_p.startswith('## '):
            doc.add_heading(clean_p[3:].replace('**', '').replace('*', ''), level=2)
        elif clean_p.startswith('### '):
            doc.add_heading(clean_p[4:].replace('**', '').replace('*', ''), level=3)
        elif clean_p.startswith('- ') or clean_p.startswith('* '):
            p_obj = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p_obj, clean_p[2:])
        else:
            p_obj = doc.add_paragraph()
            _add_formatted_text(p_obj, clean_p)
            
    doc.save(filepath)
    return filepath

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # strip single asterisks
            clean_part = part.replace('*', '')
            if clean_part:
                paragraph.add_run(clean_part)

def generate_pdf(title, content, output_dir="documents"):
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{title.replace(' ', '_')}.pdf"
    filepath = os.path.join(output_dir, filename)
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=16, style='B')
    pdf.cell(0, 10, text=title, new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(10)
    
    paragraphs = content.split('\n')
    for p in paragraphs:
        clean_p = p.strip()
        if not clean_p:
            pdf.ln(3)
            continue
            
        if clean_p.startswith('# '):
            pdf.set_font("helvetica", size=14, style='B')
            pdf.multi_cell(0, 7, text=clean_p[2:].replace('**', '').replace('*', ''))
            pdf.ln(3)
        elif clean_p.startswith('## '):
            pdf.set_font("helvetica", size=12, style='B')
            pdf.multi_cell(0, 7, text=clean_p[3:].replace('**', '').replace('*', ''))
            pdf.ln(3)
        elif clean_p.startswith('- ') or clean_p.startswith('* '):
            pdf.set_font("helvetica", size=11)
            # FPDF2 basic text, remove bold markers for simplicity in PDF
            clean_text = clean_p[2:].replace('**', '').replace('*', '')
            pdf.multi_cell(0, 7, text=f"• {clean_text}")
        else:
            pdf.set_font("helvetica", size=11)
            clean_text = clean_p.replace('**', '').replace('*', '')
            pdf.multi_cell(0, 7, text=clean_text)
            
    pdf.output(filepath)
    return filepath
